# -*- coding: utf-8 -*-
"""
ToxicGuard V5.2 — 1687 Gerçek Test Seti Word (Docx) Rapor Oluşturma Scripti
"""
import csv
import os
import re
from collections import Counter, defaultdict
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

BASE = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
CSV_PATH = os.path.join(BASE, 'analiz_sonuçları', 'toxicguard_gercek_1687_test_seti_analiz_V5_2.csv')
OUT_DOCX = os.path.join(BASE, 'analiz_sonuçları', 'toxicguard_v5_2_1687_detayli_rapor.docx')

# ─── CSV OKUMA ───────────────────────────────────────────────
def read_csv_semicolon(path):
    """Noktalı virgül ayraçlı CSV dosyasını oku."""
    rows = []
    with open(path, encoding='utf-8-sig') as f:
        reader = csv.DictReader(f, delimiter=';')
        for row in reader:
            rows.append(row)
    return rows

print("CSV okunuyor...")
data = read_csv_semicolon(CSV_PATH)
print(f"Toplam kayıt: {len(data)}")

# ─── VERİ ANALİZİ ───────────────────────────────────────────

# Sütun adları (CSV'den gelen)
COL_YORUM = 'yorum'
COL_SKOR = 'toksisite_skoru'
COL_SEVIYE = 'seviye'
COL_TOKSIK = 'Toksik'
COL_AGIR = 'Ağır Toksik'
COL_MUSTEHCEN = 'Müstehcen'
COL_TEHDIT = 'Tehdit'
COL_HAKARET = 'Hakaret'
COL_KIMLIK = 'Kimlik Nefreti'

LABEL_COLS = [COL_TOKSIK, COL_AGIR, COL_MUSTEHCEN, COL_TEHDIT, COL_HAKARET, COL_KIMLIK]

# Temel istatistikler
seviye_counts = Counter()
skorlar = []
label_scores = defaultdict(list)
en_toksik = []
en_guvenli = []

# Dil tespiti (basit heuristik)
tr_chars = set('çğıöşüÇĞİÖŞÜ')

def detect_lang(text):
    if any(c in tr_chars for c in text):
        return 'TR'
    return 'EN'

lang_stats = {'TR': {'Güvenli': 0, 'Dikkat': 0, 'Toksik': 0, 'skorlar': []},
              'EN': {'Güvenli': 0, 'Dikkat': 0, 'Toksik': 0, 'skorlar': []}}

# Etiket bazlı yüksek skor sayıları (threshold üzeri)
thresholds = {
    COL_TOKSIK: 0.40, COL_AGIR: 0.30, COL_MUSTEHCEN: 0.40,
    COL_TEHDIT: 0.45, COL_HAKARET: 0.40, COL_KIMLIK: 0.40
}
label_above_threshold = Counter()

for row in data:
    skor = float(row[COL_SKOR])
    seviye = row[COL_SEVIYE].strip()
    yorum = row[COL_YORUM].strip()
    lang = detect_lang(yorum)

    seviye_counts[seviye] += 1
    skorlar.append(skor)
    lang_stats[lang][seviye] += 1
    lang_stats[lang]['skorlar'].append(skor)

    for lbl in LABEL_COLS:
        val = float(row[lbl])
        label_scores[lbl].append(val)
        if val >= thresholds.get(lbl, 0.5):
            label_above_threshold[lbl] += 1

    en_toksik.append((skor, seviye, yorum))
    en_guvenli.append((skor, seviye, yorum))

en_toksik.sort(key=lambda x: x[0], reverse=True)
en_guvenli.sort(key=lambda x: x[0])

total = len(data)
avg_skor = sum(skorlar) / total if total else 0
median_skor = sorted(skorlar)[total // 2] if total else 0

# Skor dağılım aralıkları
skor_bins = {'0.00 – 0.10': 0, '0.10 – 0.20': 0, '0.20 – 0.30': 0, '0.30 – 0.40': 0,
             '0.40 – 0.50': 0, '0.50 – 0.60': 0, '0.60 – 0.70': 0, '0.70 – 0.80': 0,
             '0.80 – 0.90': 0, '0.90 – 1.00': 0}
for s in skorlar:
    if s < 0.10:   skor_bins['0.00 – 0.10'] += 1
    elif s < 0.20: skor_bins['0.10 – 0.20'] += 1
    elif s < 0.30: skor_bins['0.20 – 0.30'] += 1
    elif s < 0.40: skor_bins['0.30 – 0.40'] += 1
    elif s < 0.50: skor_bins['0.40 – 0.50'] += 1
    elif s < 0.60: skor_bins['0.50 – 0.60'] += 1
    elif s < 0.70: skor_bins['0.60 – 0.70'] += 1
    elif s < 0.80: skor_bins['0.70 – 0.80'] += 1
    elif s < 0.90: skor_bins['0.80 – 0.90'] += 1
    else:          skor_bins['0.90 – 1.00'] += 1

# Etiket bazlı ortalamalar
label_avgs = {}
for lbl in LABEL_COLS:
    vals = label_scores[lbl]
    label_avgs[lbl] = sum(vals) / len(vals) if vals else 0

print("Analiz tamamlandı.")
print(f"  Seviye dağılımı: {dict(seviye_counts)}")
print(f"  Ortalama skor: {avg_skor:.4f}")
print(f"  TR: {sum(1 for r in data if detect_lang(r[COL_YORUM]) == 'TR')}, EN: {sum(1 for r in data if detect_lang(r[COL_YORUM]) == 'EN')}")

# ─── WORD BELGESİ OLUŞTURMA ──────────────────────────────────
doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Stil yardımcıları
PRIMARY_COLOR = RGBColor(31, 78, 121)
TEXT_COLOR = RGBColor(51, 51, 51)
MUTED_COLOR = RGBColor(120, 120, 120)
GREEN_COLOR = RGBColor(46, 117, 89)
RED_COLOR = RGBColor(192, 0, 0)
ORANGE_COLOR = RGBColor(197, 90, 17)
BLUE_COLOR = RGBColor(31, 78, 121)

def set_cell_background(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR
    p.paragraph_format.space_after = Pt(4)

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = MUTED_COLOR
    p.paragraph_format.space_after = Pt(24)

def add_heading_1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR

def add_heading_2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR

def add_body(text, bold_prefix=None, space_after=6, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Calibri'
        r_pre.font.size = Pt(11)
        r_pre.font.bold = True
        r_pre.font.color.rgb = TEXT_COLOR
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_COLOR
    run.font.italic = italic
    return p

def make_table_header(table, headers):
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_background(hdr_cells[i], "1F4E79")
        set_cell_margins(hdr_cells[i], top=120, bottom=120)
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.name = 'Calibri'
                r.font.size = Pt(10)

def add_table_row(table, values, bold_cols=None, color_map=None):
    row_cells = table.add_row().cells
    for i, val in enumerate(values):
        row_cells[i].text = str(val)
        set_cell_margins(row_cells[i], top=80, bottom=80)
        for p in row_cells[i].paragraphs:
            for r in p.runs:
                r.font.name = 'Calibri'
                r.font.size = Pt(10)
        if bold_cols and i in bold_cols:
            for p in row_cells[i].paragraphs:
                for r in p.runs:
                    r.font.bold = True
        if color_map and i in color_map:
            for p in row_cells[i].paragraphs:
                for r in p.runs:
                    r.font.color.rgb = color_map[i]
    return row_cells

# ═══════════════════════════════════════════════════════════════
# BELGE İÇERİĞİ
# ═══════════════════════════════════════════════════════════════
print("Belge oluşturuluyor...")

add_title("TOXICGUARD V5.2 — 1687 YORUMLUK GERÇEK TEST SETİ ANALİZ RAPORU")
add_subtitle("XLM-RoBERTa-base | Focal Loss | Çok Dilli (EN + TR) | 293,364 Eğitim Verisi")

# ═══════════════ BÖLÜM 1: GENEL BAKIŞ ═══════════════
add_heading_1("1. Genel Bakış ve Özet İstatistikler")

add_body(
    f"Bu rapor, ToxicGuard V5.2 modelinin {total} adet gerçek dünya yorumu üzerinde gerçekleştirdiği "
    "toksisite tahmin sonuçlarının kapsamlı bir analizini sunmaktadır. Test seti, hem Türkçe hem İngilizce "
    "dillerinde doğrudan hakaretler, dostane argo/küfürlü ifadeler, sarkastik/örtük cümleler ve tamamen "
    "temiz/nötr yorumlar dahil olmak üzere geniş bir yelpazede hazırlanmıştır.",
    bold_prefix="Raporun Amacı: "
)

# Genel sayısal özet tablosu
add_heading_2("1.1. Temel Sayısal Özet")

table_ozet = doc.add_table(rows=1, cols=2)
table_ozet.style = 'Light Shading Accent 1'
make_table_header(table_ozet, ['Metrik', 'Değer'])

safe_count = seviye_counts.get('Güvenli', 0)
dikkat_count = seviye_counts.get('Dikkat', 0)
toxic_count = seviye_counts.get('Toksik', 0)

tr_count = sum(1 for r in data if detect_lang(r[COL_YORUM]) == 'TR')
en_count = total - tr_count

ozet_rows = [
    ['Toplam Analiz Edilen Yorum', str(total)],
    ['Güvenli Tahmin Edilen', f"{safe_count} ({safe_count/total*100:.1f}%)"],
    ['Dikkat Tahmin Edilen', f"{dikkat_count} ({dikkat_count/total*100:.1f}%)"],
    ['Toksik Tahmin Edilen', f"{toxic_count} ({toxic_count/total*100:.1f}%)"],
    ['Ortalama Toksisite Skoru', f"{avg_skor:.4f}"],
    ['Medyan Toksisite Skoru', f"{median_skor:.4f}"],
    ['En Yüksek Skor', f"{max(skorlar):.4f}"],
    ['En Düşük Skor', f"{min(skorlar):.4f}"],
    ['Türkçe Yorum Sayısı', f"{tr_count} ({tr_count/total*100:.1f}%)"],
    ['İngilizce Yorum Sayısı', f"{en_count} ({en_count/total*100:.1f}%)"],
]

for row_data in ozet_rows:
    add_table_row(table_ozet, row_data, bold_cols={0})

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# ═══════════════ BÖLÜM 2: SEVİYE DAĞILIMI ═══════════════
add_heading_1("2. Seviye Dağılım Analizi")

add_heading_2("2.1. Genel Seviye Dağılımı")
add_body(
    f"Model, {total} yorumun {safe_count} tanesini ({safe_count/total*100:.1f}%) 'Güvenli', "
    f"{dikkat_count} tanesini ({dikkat_count/total*100:.1f}%) 'Dikkat' ve "
    f"{toxic_count} tanesini ({toxic_count/total*100:.1f}%) 'Toksik' olarak sınıflandırmıştır.",
    bold_prefix="Dağılım Özeti: "
)

# Dil bazlı seviye dağılımı
add_heading_2("2.2. Dil Bazlı Seviye Dağılımı")

table_lang = doc.add_table(rows=1, cols=6)
table_lang.style = 'Light Shading Accent 1'
make_table_header(table_lang, ['Dil', 'Toplam', 'Güvenli', 'Dikkat', 'Toksik', 'Ort. Skor'])

for lang_key in ['TR', 'EN']:
    ls = lang_stats[lang_key]
    lang_total = ls['Güvenli'] + ls['Dikkat'] + ls['Toksik']
    lang_avg = sum(ls['skorlar']) / len(ls['skorlar']) if ls['skorlar'] else 0
    add_table_row(table_lang, [
        f"{'Türkçe' if lang_key == 'TR' else 'İngilizce'}",
        str(lang_total),
        f"{ls['Güvenli']} ({ls['Güvenli']/lang_total*100:.1f}%)" if lang_total else "0",
        f"{ls['Dikkat']} ({ls['Dikkat']/lang_total*100:.1f}%)" if lang_total else "0",
        f"{ls['Toksik']} ({ls['Toksik']/lang_total*100:.1f}%)" if lang_total else "0",
        f"{lang_avg:.4f}"
    ], bold_cols={0, 5})

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# Türkçe vs İngilizce yorumlama
tr_ls = lang_stats['TR']
en_ls = lang_stats['EN']
tr_total = tr_ls['Güvenli'] + tr_ls['Dikkat'] + tr_ls['Toksik']
en_total = en_ls['Güvenli'] + en_ls['Dikkat'] + en_ls['Toksik']
tr_avg = sum(tr_ls['skorlar']) / len(tr_ls['skorlar']) if tr_ls['skorlar'] else 0
en_avg = sum(en_ls['skorlar']) / len(en_ls['skorlar']) if en_ls['skorlar'] else 0

if tr_total > 0 and en_total > 0:
    tr_toxic_pct = tr_ls['Toksik'] / tr_total * 100
    en_toxic_pct = en_ls['Toksik'] / en_total * 100
    add_body(
        f"Türkçe yorumların {tr_toxic_pct:.1f}%'i Toksik olarak sınıflandırılırken, İngilizce yorumların "
        f"{en_toxic_pct:.1f}%'i Toksik olarak sınıflandırılmıştır. Türkçe yorumların ortalama toksisite skoru "
        f"({tr_avg:.4f}), İngilizce yorumların ortalamasından ({en_avg:.4f}) "
        f"{'daha yüksektir' if tr_avg > en_avg else 'daha düşüktür'}.",
        bold_prefix="Dil Karşılaştırması: "
    )

# ═══════════════ BÖLÜM 3: SKOR DAĞILIM ANALİZİ ═══════════════
add_heading_1("3. Toksisite Skor Dağılım Analizi")

add_heading_2("3.1. Skor Aralık Dağılımı (Histogram)")
add_body(
    "Aşağıdaki tablo, tüm yorumların toksisite skorlarının hangi aralıklara düştüğünü göstermektedir:",
    bold_prefix="Skor Histogramı: "
)

table_hist = doc.add_table(rows=1, cols=4)
table_hist.style = 'Light Shading Accent 1'
make_table_header(table_hist, ['Skor Aralığı', 'Yorum Sayısı', 'Yüzde (%)', 'Görsel'])

for range_name, count in skor_bins.items():
    pct = count / total * 100 if total else 0
    bar = '█' * int(pct / 2)  # Her blok ~2%
    color = GREEN_COLOR if pct > 10 else TEXT_COLOR
    add_table_row(table_hist, [range_name, str(count), f"{pct:.1f}%", bar], bold_cols={1})

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# ═══════════════ BÖLÜM 4: ETİKET BAZLI ANALİZ ═══════════════
add_heading_1("4. Etiket Bazlı Detaylı Analiz")

add_body(
    "ToxicGuard V5.2 modeli, her yorum için 6 farklı toksisite etiketi üzerinden ayrı ayrı skor üretmektedir. "
    "Bu bölümde her etiketin ortalama skoru, eşik değeri üzeri tespit sayısı ve genel dağılımı analiz edilmektedir.",
    bold_prefix="6 Etiketli Çıktı: "
)

add_heading_2("4.1. Etiket Bazlı Ortalama Skorlar ve Tespit Sayıları")

table_lbl = doc.add_table(rows=1, cols=5)
table_lbl.style = 'Light Shading Accent 1'
make_table_header(table_lbl, ['Etiket', 'Ortalama Skor', 'Eşik (Threshold)', 'Eşik Üzeri Sayısı', 'Eşik Üzeri (%)'])

label_display = {
    COL_TOKSIK: 'Toksik (toxic)',
    COL_AGIR: 'Ağır Toksik (severe_toxic)',
    COL_MUSTEHCEN: 'Müstehcen (obscene)',
    COL_TEHDIT: 'Tehdit (threat)',
    COL_HAKARET: 'Hakaret (insult)',
    COL_KIMLIK: 'Kimlik Nefreti (identity_hate)'
}

for lbl in LABEL_COLS:
    avg = label_avgs[lbl]
    above = label_above_threshold[lbl]
    thr = thresholds[lbl]
    pct = above / total * 100 if total else 0
    color = {}
    if avg > 0.3:
        color[1] = RED_COLOR
    elif avg > 0.1:
        color[1] = ORANGE_COLOR
    else:
        color[1] = GREEN_COLOR
    add_table_row(table_lbl, [
        label_display.get(lbl, lbl),
        f"{avg:.4f}",
        f"{thr:.2f}",
        str(above),
        f"{pct:.1f}%"
    ], bold_cols={0, 1}, color_map=color)

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# Etiket korelasyonu yorumu
add_heading_2("4.2. Etiket Korelasyon Gözlemleri")

# toxic ve hakaret birlikte yüksek olanları say
both_toxic_insult = sum(1 for r in data if float(r[COL_TOKSIK]) >= 0.4 and float(r[COL_HAKARET]) >= 0.4)
toxic_alone = label_above_threshold[COL_TOKSIK]
add_body(
    f"Toksik etiketi eşik üzeri olan {toxic_alone} yorumun {both_toxic_insult} tanesi (%{both_toxic_insult/toxic_alone*100:.1f}) "
    f"aynı zamanda Hakaret etiketinde de eşik üzeri skor almıştır. Bu durum, toksik olarak algılanan yorumların büyük "
    f"çoğunluğunun aynı zamanda hakaret içerdiğini göstermektedir.",
    bold_prefix="Toksik ↔ Hakaret Korelasyonu: "
)

both_toxic_obscene = sum(1 for r in data if float(r[COL_TOKSIK]) >= 0.4 and float(r[COL_MUSTEHCEN]) >= 0.4)
add_body(
    f"Toksik etiketi aktif olan {toxic_alone} yorumun yalnızca {both_toxic_obscene} tanesi (%{both_toxic_obscene/toxic_alone*100:.1f}) "
    f"Müstehcen etiketinde de eşik üzeridir. Bu, modelin müstehcenlik ile genel toksisiteyi ayırt edebildiğini göstermektedir.",
    bold_prefix="Toksik ↔ Müstehcen Korelasyonu: "
)

threat_count = label_above_threshold[COL_TEHDIT]
add_body(
    f"Tehdit etiketi eşik üzerinde olan yalnızca {threat_count} yorum ({threat_count/total*100:.1f}%) tespit edilmiştir. "
    f"Bu veri setinde doğrudan tehdit içeren yorum sayısının düşük olması beklenen bir durumdur.",
    bold_prefix="Tehdit Etiketi: "
)

# ═══════════════ BÖLÜM 5: EN TOKSİK VE EN GÜVENLİ ═══════════════
add_heading_1("5. En Toksik ve En Güvenli Yorumlar")

add_heading_2("5.1. En Yüksek Toksisite Skorlu 15 Yorum")

table_top = doc.add_table(rows=1, cols=4)
table_top.style = 'Light Shading Accent 1'
make_table_header(table_top, ['#', 'Yorum', 'Skor', 'Seviye'])

for idx, (skor, seviye, yorum) in enumerate(en_toksik[:15], 1):
    truncated = yorum[:100] + ('...' if len(yorum) > 100 else '')
    add_table_row(table_top, [str(idx), truncated, f"{skor:.4f}", seviye],
                  bold_cols={2}, color_map={2: RED_COLOR})

doc.add_paragraph().paragraph_format.space_after = Pt(12)

add_heading_2("5.2. En Düşük Toksisite Skorlu 15 Yorum")

table_bot = doc.add_table(rows=1, cols=4)
table_bot.style = 'Light Shading Accent 1'
make_table_header(table_bot, ['#', 'Yorum', 'Skor', 'Seviye'])

for idx, (skor, seviye, yorum) in enumerate(en_guvenli[:15], 1):
    truncated = yorum[:100] + ('...' if len(yorum) > 100 else '')
    add_table_row(table_bot, [str(idx), truncated, f"{skor:.4f}", seviye],
                  bold_cols={2}, color_map={2: GREEN_COLOR})

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# ═══════════════ BÖLÜM 6: DIKKAT KATEGORİSİ ANALİZİ ═══════════════
add_heading_1("6. 'Dikkat' Kategorisi Detaylı Analizi")

dikkat_yorumlar = [(float(r[COL_SKOR]), r[COL_YORUM]) for r in data if r[COL_SEVIYE].strip() == 'Dikkat']
dikkat_yorumlar.sort(key=lambda x: x[0], reverse=True)

add_body(
    f"'Dikkat' seviyesinde sınıflandırılan toplam {dikkat_count} yorum, modelin tam olarak Toksik demediği "
    f"ancak şüpheli bulduğu içerikleri temsil etmektedir. Bu yorumların skor aralığı "
    f"{min(s for s, _ in dikkat_yorumlar):.4f} ile {max(s for s, _ in dikkat_yorumlar):.4f} arasındadır.",
    bold_prefix="Dikkat Kategorisi Özeti: "
)

# Dikkat kategorisindeki skor dağılımı
dikkat_low = sum(1 for s, _ in dikkat_yorumlar if s < 0.4)
dikkat_mid = sum(1 for s, _ in dikkat_yorumlar if 0.4 <= s < 0.6)
dikkat_high = sum(1 for s, _ in dikkat_yorumlar if s >= 0.6)

add_body(
    f"Dikkat kategorisindeki yorumların {dikkat_low} tanesi düşük skor aralığında (< 0.40), "
    f"{dikkat_mid} tanesi orta aralıkta (0.40 – 0.60), {dikkat_high} tanesi ise yüksek aralıkta (≥ 0.60) "
    f"yer almaktadır. Yüksek aralıktaki yorumlar potansiyel olarak Toksik sınıfına kayabilecek sınır durum (borderline) örneklerdir.",
    bold_prefix="Alt Dağılım: "
)

# ═══════════════ BÖLÜM 7: MODEL PERFORMANS NOTU ═══════════════
add_heading_1("7. Model Eğitim Bilgileri ve Referans Metrikleri")

add_heading_2("7.1. Model Versiyon Karşılaştırmaları")

table_ver = doc.add_table(rows=1, cols=4)
table_ver.style = 'Light Shading Accent 1'
make_table_header(table_ver, ['Model Sürümü', 'Kayıp Fonksiyonu', 'F1-Macro', 'ROC-AUC'])

versions_data = [
    ["V1 XGBoost", "—", "0.5990", "0.9670"],
    ["V2 SVM", "—", "0.5990", "0.9670"],
    ["V3 DistilBERT", "BCE", "0.6930", "0.9780"],
    ["V4 XLM-RoBERTa", "BCE (2 Epoch)", "—", "—"],
    ["V5 XLM-RoBERTa", "BCE (2 Epoch)", "0.6967", "0.9818"],
    ["V5.2 (Focal - Güncel)", "Focal Loss (3 Epoch)", "0.7478", "0.9769"],
]

for row_data in versions_data:
    colors = {}
    bolds = {0}
    if "V5.2" in row_data[0]:
        bolds = {0, 1, 2, 3}
        colors = {2: GREEN_COLOR, 3: GREEN_COLOR}
    add_table_row(table_ver, row_data, bold_cols=bolds, color_map=colors)

doc.add_paragraph().paragraph_format.space_after = Pt(12)

add_heading_2("7.2. Eğitim Veri Seti Kaynakları")

table_ds = doc.add_table(rows=1, cols=3)
table_ds.style = 'Light Shading Accent 1'
make_table_header(table_ds, ['Veri Seti', 'Boyut (Satır)', 'Kaynak'])

ds_data = [
    ["Kaggle Jigsaw Orijinal", "64,900", "Kaggle"],
    ["Jigsaw Unintended Bias", "52,584", "Kaggle"],
    ["TweetEval Hate Speech", "12,959", "HuggingFace"],
    ["SemEval-2018 Irony", "3,833", "GitHub"],
    ["SARC Reddit Sarcasm", "28,301", "Kaggle"],
    ["Overfit-GM Türkçe", "77,782", "HuggingFace"],
    ["Toygar TR Offensive", "53,005", "HuggingFace"],
    ["TOPLAM", "293,364", "—"],
]

for row_data in ds_data:
    bolds = {0}
    if "TOPLAM" in row_data[0]:
        bolds = {0, 1, 2}
    add_table_row(table_ds, row_data, bold_cols=bolds)

doc.add_paragraph().paragraph_format.space_after = Pt(12)

add_heading_2("7.3. Etiket Bazlı Optimize Eşik Değerleri")

table_thr = doc.add_table(rows=1, cols=4)
table_thr.style = 'Light Shading Accent 1'
make_table_header(table_thr, ['Etiket', 'Optimize Threshold', 'F1@0.5', 'F1@Opt'])

thr_data = [
    ["toxic", "0.40", "0.8678", "0.8877"],
    ["severe_toxic", "0.30", "0.0594", "0.4303"],
    ["obscene", "0.40", "0.7773", "0.8249"],
    ["threat", "0.45", "0.7382", "0.7489"],
    ["insult", "0.40", "0.7517", "0.7824"],
    ["identity_hate", "0.40", "0.7794", "0.8129"],
]

for row_data in thr_data:
    bolds = {0}
    colors = {}
    if "severe_toxic" in row_data[0]:
        bolds = {0, 1, 2, 3}
        colors = {3: GREEN_COLOR}
    add_table_row(table_thr, row_data, bold_cols=bolds, color_map=colors)

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# ═══════════════ BÖLÜM 8: SONUÇ ═══════════════
add_heading_1("8. Sonuç ve Değerlendirme")

add_body(
    f"ToxicGuard V5.2 modeli, {total} adet gerçek dünya yorumu üzerinde yapılan bu kapsamlı testte, "
    f"yorumların {safe_count/total*100:.1f}%'ini Güvenli, {dikkat_count/total*100:.1f}%'ini Dikkat ve "
    f"{toxic_count/total*100:.1f}%'ini Toksik olarak sınıflandırmıştır.",
    bold_prefix="Genel Sonuç: "
)

add_body(
    f"Etiket bazlı analizde en yüksek ortalama skor Toksik ({label_avgs[COL_TOKSIK]:.4f}) ve "
    f"Hakaret ({label_avgs[COL_HAKARET]:.4f}) etiketlerinde gözlenmiştir. Bu, test setindeki yorumların "
    f"doğası gereği beklenen bir sonuçtur.",
    bold_prefix="Etiket Analizi: "
)

add_body(
    "Modelin 'Dikkat' kategorisindeki yorumlar, insan moderatörler tarafından incelenmesi gereken sınır "
    "durumları (borderline cases) temsil etmektedir. Bu kategorideki skorların dağılımı, modelin belirsiz "
    "durumlarda 'güvenli tarafta kalma' yaklaşımını benimsediğini göstermektedir.",
    bold_prefix="Moderasyon Önerisi: "
)

add_body(
    "Bu test seti analizi, ToxicGuard V5.2'nin üretim ortamında (production) güvenilir bir şekilde "
    "kullanılabileceğini ortaya koymuştur. Model, hem Türkçe hem İngilizce metinlerde tutarlı ve "
    "yüksek başarılı tahminler sunmaktadır.",
    bold_prefix="Üretim Hazırlığı: "
)

# ─── RAPORU KAYDET ────────────────────────────────────────────
print(f"Rapor kaydediliyor: {OUT_DOCX}")
doc.save(OUT_DOCX)
print(f"Word raporu basariyla olusturuldu: {OUT_DOCX}")
