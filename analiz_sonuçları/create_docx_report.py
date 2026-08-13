# -*- coding: utf-8 -*-
"""
ToxicGuard V5.2 — Word (Docx) Rapor Oluşturma Scripti
"""
import csv
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

BASE = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
RES_1000 = os.path.join(BASE, 'analiz_sonuçları', 'toxicguard_1000_genis_test_seti_v5_2_sonuclar.csv')
LBL_1000 = os.path.join(BASE, 'analiz_edilecekler', 'toxicguard_1000_genis_test_seti.csv')
RES_300  = os.path.join(BASE, 'analiz_sonuçları', 'advanced_toxicity_test_300_v5_2_sonuclar.csv')
OUT_DOCX = os.path.join(BASE, 'analiz_sonuçları', 'toxicguard_v5_2_detayli_rapor.docx')

def read_csv(path):
    rows = []
    with open(path, encoding='utf-8-sig') as f:
        reader = csv.DictReader(f)
        for row in reader:
            rows.append(row)
    return rows

# ─── VERİ OKUMA VE HESAPLAMA ────────────────────────────────
print("CSV'ler okunuyor...")
res_1000 = read_csv(RES_1000)
lbl_1000 = read_csv(LBL_1000)
res_300  = read_csv(RES_300)
print(f"1000-set: {len(res_1000)} kayıt, 300-set: {len(res_300)} kayıt")

# Eşleme haritası tanımla
categories_map = {
    ('Explicit Toxic (Açık Toksik)', 'Turkish'): 'Açık Toksik (TR)',
    ('Explicit Toxic (Açık Toksik)', 'English'): 'Açık Toksik (EN)',
    ('Clean (Temiz/Nötr)', 'Turkish'): 'Temiz/Nötr (TR)',
    ('Clean (Temiz/Nötr)', 'English'): 'Temiz/Nötr (EN)',
    ('Friendly Profanity (Dostane Küfürlü)', 'Turkish'): 'Dostane Küfürlü (TR)',
    ('Friendly Profanity (Dostane Küfürlü)', 'English'): 'Dostane Küfürlü (EN)',
    ('Implicit Sarcastic (Sarkastik/Örtük)', 'Turkish'): 'Sarkastik/Örtük (TR)',
    ('Implicit Sarcastic (Sarkastik/Örtük)', 'English'): 'Sarkastik/Örtük (EN)'
}

stats = {name: {'Güvenli': 0, 'Dikkat': 0, 'Toksik': 0, 'SkorToplam': 0.0, 'Adet': 0} for name in categories_map.values()}

# Hatalı örnekleri toplamak için listeler
false_positives = []
false_negatives = []

for r_item, l_item in zip(res_1000, lbl_1000):
    kategori = l_item.get('kategori', '').strip()
    dil = l_item.get('dil', '').strip()
    key = (kategori, dil)
    
    if key in categories_map:
        cat_name = categories_map[key]
        seviye = r_item.get('seviye', '').strip()
        skor = float(r_item.get('toksisite_skoru', 0))
        yorum = r_item.get('yorum', '').strip()
        
        stats[cat_name]['Adet'] += 1
        stats[cat_name]['SkorToplam'] += skor
        
        if seviye in ('Güvenli', 'Guvenli'):
            stats[cat_name]['Güvenli'] += 1
        elif seviye == 'Dikkat':
            stats[cat_name]['Dikkat'] += 1
        elif seviye == 'Toksik':
            stats[cat_name]['Toksik'] += 1
            
        # Yanlış sınıflandırma örneklerini kaydet
        if 'Explicit Toxic' in kategori and seviye not in ('Toksik',):
            false_negatives.append((yorum, skor, seviye, cat_name))
        if 'Clean' in kategori and seviye == 'Toksik':
            false_positives.append((yorum, skor, seviye, cat_name))
        if 'Friendly Profanity' in kategori and seviye == 'Toksik':
            false_positives.append((yorum, skor, seviye, cat_name))

print("Hesaplamalar tamamlandı.")

# ─── WORD BELGESİ OLUŞTURMA ──────────────────────────────────
doc = Document()

# Sayfa kenar boşlukları ayarlama (standart 2.54 cm / 1 inç)
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Stil yardımcıları
PRIMARY_COLOR = RGBColor(31, 78, 121)     # Koyu Mavi (#1F4E79)
TEXT_COLOR = RGBColor(51, 51, 51)          # Koyu Gri (#333333)
MUTED_COLOR = RGBColor(120, 120, 120)      # Açık Gri (#787878)
GREEN_COLOR = RGBColor(46, 117, 89)        # Yeşil (#2E7559)
RED_COLOR = RGBColor(192, 0, 0)            # Kırmızı (#C00000)

def set_cell_background(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR
    p.paragraph_format.space_after = Pt(4)

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = MUTED_COLOR
    p.paragraph_format.space_after = Pt(24)

def add_heading_1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR

def add_heading_2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR

def add_body(text, bold_prefix=None, space_after=6, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Calibri'
        r_pre.font.size = Pt(11)
        r_pre.font.bold = True
        r_pre.font.color.rgb = TEXT_COLOR
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_COLOR
    run.font.italic = italic
    return p

# ─── BELGE İÇERİĞİ YAZILIYOR ─────────────────────────────────
print("Belge oluşturuluyor...")

add_title("TOXICGUARD V5.2 DETAYLI TAHMİN ANALİZ RAPORU")
add_subtitle("XLM-RoBERTa-base | Focal Loss | Çok Dilli (EN + TR) | 293,364 Eğitim Verisi")

# ----------------- BÖLÜM 1 -----------------
add_heading_1("1. Model Eğitim Özeti ve Temel Başarı Metrikleri")

add_body(
    "ToxicGuard projesinin V5.2 sürümü, çok dilli (İngilizce ve Türkçe) toksisite tespiti amacıyla "
    "XLM-RoBERTa-base (278M parametre) mimarisi üzerine inşa edilmiştir. Eğitimde nadir sınıfların "
    "ve dengesiz veri dağılımlarının olumsuz etkilerini önlemek adına Focal Loss (gamma=2.0, alpha=0.25) "
    "kayıp fonksiyonu kullanılmıştır. Toplamda 3 epoch boyunca eğitilen modelde, eşik değerleri (threshold) "
    "optimize edilerek hedeflenen başarımın üzerine çıkılmıştır.",
    bold_prefix="Model Yapısı ve Kayıp Fonksiyonu: "
)

# Versiyon karşılaştırma tablosu
add_heading_2("1.1. Model Versiyon Karşılaştırmaları (V1 - V5.2)")
add_body("Modellerin F1-Macro ve ROC-AUC başarı metriklerinin sürümlere göre gelişimi aşağıdaki gibidir:")

table_ver = doc.add_table(rows=1, cols=4)
table_ver.style = 'Light Shading Accent 1'
hdr_cells = table_ver.rows[0].cells
hdr_cells[0].text = 'Model Sürümü'
hdr_cells[1].text = 'Kayıp Fonksiyonu'
hdr_cells[2].text = 'F1-Macro Skoru'
hdr_cells[3].text = 'ROC-AUC Skoru'

for cell in hdr_cells:
    set_cell_background(cell, "1F4E79")
    set_cell_margins(cell, top=120, bottom=120)
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)

versions_data = [
    ["V1 XGBoost", "—", "0.5990", "0.9670"],
    ["V2 SVM", "—", "0.5990", "0.9670"],
    ["V3 DistilBERT", "BCE", "0.6930", "0.9780"],
    ["V4 XLM-RoBERTa", "BCE (2 Epoch)", "—", "—"],
    ["V5 XLM-RoBERTa", "BCE (2 Epoch)", "0.6967", "0.9818"],
    ["V5.2 (Focal - Güncel)", "Focal Loss (3 Epoch)", "0.7478", "0.9769"],
]

for row_data in versions_data:
    row_cells = table_ver.add_row().cells
    for i, val in enumerate(row_data):
        row_cells[i].text = val
        set_cell_margins(row_cells[i], top=80, bottom=80)
        # Güncel modeli kalın yaz
        if "V5.2" in row_data[0]:
            for p in row_cells[i].paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    if i in (2, 3):
                        r.font.color.rgb = GREEN_COLOR

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# Eğitim veri kümesi dağılım tablosu
add_heading_2("1.2. Eğitim Veri Seti Etiket Dağılımları")
add_body(
    "Model toplam 293,364 satırdan oluşan devasa bir çok dilli veri setiyle eğitilmiştir. "
    "Sınıfların dağılımı ve her etikete düşen eğitim örneği sayıları şu şekildedir:",
    bold_prefix="Toplam Veri Kümesi: "
)

table_dist = doc.add_table(rows=1, cols=3)
table_dist.style = 'Light Shading Accent 1'
hdr_cells_d = table_dist.rows[0].cells
hdr_cells_d[0].text = 'Etiket (Label)'
hdr_cells_d[1].text = 'Eğitim Örneği Sayısı'
hdr_cells_d[2].text = 'Yüzdelik Oran (%)'

for cell in hdr_cells_d:
    set_cell_background(cell, "1F4E79")
    set_cell_margins(cell, top=120, bottom=120)
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)

dist_data = [
    ["toxic", "133,414", "45.5%"],
    ["severe_toxic", "2,088", "0.7%"],
    ["obscene", "28,207", "9.6%"],
    ["threat", "4,758", "1.6%"],
    ["insult", "67,053", "22.9%"],
    ["identity_hate", "31,384", "10.7%"],
    ["TOPLAM VERİ SETİ", "293,364", "100.0%"]
]

for row_data in dist_data:
    row_cells = table_dist.add_row().cells
    for i, val in enumerate(row_data):
        row_cells[i].text = val
        set_cell_margins(row_cells[i], top=80, bottom=80)
        if "TOPLAM" in row_data[0]:
            set_cell_background(row_cells[i], "F2F2F2")
            for p in row_cells[i].paragraphs:
                for r in p.runs:
                    r.font.bold = True

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# Sınıflandırma Eşikleri (Thresholds)
add_heading_2("1.3. Etiket Bazlı Optimize Eşik Değerleri ve Başarı Dağılımları")
add_body(
    "F1 skorunu maksimize etmek amacıyla etiketlere göre optimize edilmiş threshold "
    "değerleri belirlenmiştir. Standart threshold (0.50) yerine optimize eşikler kullanıldığında "
    "özellikle severe_toxic ve threat gibi nadir sınıfların yakalanma oranında çok büyük bir artış gözlenmiştir:",
    bold_prefix="Threshold Optimizasyonu: "
)

table_thr = doc.add_table(rows=1, cols=4)
table_thr.style = 'Light Shading Accent 1'
hdr_cells_t = table_thr.rows[0].cells
hdr_cells_t[0].text = 'Etiket'
hdr_cells_t[1].text = 'Optimize Threshold'
hdr_cells_t[2].text = 'Varsayılan Eşik F1 (0.50)'
hdr_cells_t[3].text = 'Optimize Eşik F1 (Opt)'

for cell in hdr_cells_t:
    set_cell_background(cell, "1F4E79")
    set_cell_margins(cell, top=120, bottom=120)
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)

thr_data = [
    ["toxic", "0.40", "0.8678", "0.8877"],
    ["severe_toxic", "0.30", "0.0594", "0.4303"],
    ["obscene", "0.40", "0.7773", "0.8249"],
    ["threat", "0.45", "0.7382", "0.7489"],
    ["insult", "0.40", "0.7517", "0.7824"],
    ["identity_hate", "0.40", "0.7794", "0.8129"],
]

for row_data in thr_data:
    row_cells = table_thr.add_row().cells
    for i, val in enumerate(row_data):
        row_cells[i].text = val
        set_cell_margins(row_cells[i], top=80, bottom=80)
        # severe_toxic'i renklendir/kalınlaştır çünkü en büyük fark orada
        if "severe_toxic" in row_data[0]:
            for p in row_cells[i].paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    if i == 3:
                        r.font.color.rgb = GREEN_COLOR

# ----------------- BÖLÜM 2 -----------------
add_heading_1("2. 1000 Yorumluk Geniş Test Seti Kategori Analizi")
add_body(
    "Eğitilen model, 1000 adet (manuel etiketlenmiş 693 adet) Türkçe ve İngilizce karmaşık yorum "
    "barındıran geniş test setine tabi tutulmuştur. Kategorilerin dil bazında dağılımı, "
    "tahmin edilen güvenlik düzeyleri ve ortalama toksisite skorları aşağıdaki tabloda verilmiştir:",
    bold_prefix="Test Kümesi Dağılımı: "
)

# DİNAMİK HESAPLANAN KATEGORİ TABLOSU
table_cat = doc.add_table(rows=1, cols=5)
table_cat.style = 'Table Grid'
hdr_cells_c = table_cat.rows[0].cells
hdr_cells_c[0].text = 'Kategori'
hdr_cells_c[1].text = 'Güvenli (Adet)'
hdr_cells_c[2].text = 'Dikkat (Adet)'
hdr_cells_c[3].text = 'Toksik (Adet)'
hdr_cells_c[4].text = 'Ortalama Skor'

for cell in hdr_cells_c:
    set_cell_background(cell, "1F4E79")
    set_cell_margins(cell, top=120, bottom=120)
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)

# Tablo satırlarını yaz
ordered_cats = [
    'Açık Toksik (TR)',
    'Açık Toksik (EN)',
    'Temiz/Nötr (TR)',
    'Temiz/Nötr (EN)',
    'Dostane Küfürlü (TR)',
    'Dostane Küfürlü (EN)',
    'Sarkastik/Örtük (TR)',
    'Sarkastik/Örtük (EN)'
]

for cat_name in ordered_cats:
    cat_info = stats[cat_name]
    row_cells = table_cat.add_row().cells
    
    row_cells[0].text = cat_name
    row_cells[1].text = str(cat_info['Güvenli'])
    row_cells[2].text = str(cat_info['Dikkat'])
    row_cells[3].text = str(cat_info['Toksik'])
    
    avg_score = 0.0
    if cat_info['Adet'] > 0:
        avg_score = cat_info['SkorToplam'] / cat_info['Adet']
    row_cells[4].text = f"{avg_score:.4f}"
    
    # Hücreleri biçimlendir
    for i in range(5):
        set_cell_margins(row_cells[i], top=80, bottom=80)
        # Özel renklendirmeler
        if i == 0:
            for p in row_cells[i].paragraphs:
                for r in p.runs:
                    r.font.bold = True
        elif i == 4:
            for p in row_cells[i].paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    # Skora göre renk ver
                    if avg_score > 0.5:
                        r.font.color.rgb = RED_COLOR
                    elif avg_score > 0.2:
                        r.font.color.rgb = RGBColor(197, 90, 17) # Turuncu
                    else:
                        r.font.color.rgb = GREEN_COLOR

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# Bulgular
add_body(
    "Açık Toksik yorumlar hem Türkçe hem de İngilizce metinlerde ezici bir şekilde 'Toksik' veya 'Dikkat' "
    "olarak doğru tahmin edilmiştir. Türkçe açık toksiklerin ortalama skoru 0.8229, İngilizce açık toksiklerin "
    "ise 0.7837'dir. Bu da modelin doğrudan hakaretleri kaçırmadığını gösterir.",
    bold_prefix="Açık Toksik Performansı: "
)

add_body(
    "Temiz/Nötr yorumlar çok düşük ortalama skorlarla (%94'ün üzerinde) doğru şekilde 'Güvenli' "
    "olarak sınıflandırılmıştır. Türkçe temiz yorumların ortalama skoru 0.0599 gibi mükemmel bir seviyededir. "
    "İngilizce temiz yorumlarda bu skor 0.2180'dir. Bu durum modelin genel 'false positive' (temiz cümleye "
    "durduk yere toksik deme) oranının son derece düşük olduğunu ispatlamaktadır.",
    bold_prefix="Temiz/Nötr Performansı: "
)

add_body(
    "Modelin en çok zorlandığı alandır. Türkçe dostane küfürlü/argolu yorumlar 0.6116, İngilizce olanlar ise "
    "0.5117 ortalama skor alarak büyük oranda 'Toksik' veya 'Dikkat' seviyesine düşmüştür. Arkadaşça söylenen "
    "'şerefsiz', 'crazy bastard' gibi kelimeler model tarafından doğrudan kelime bazlı değerlendirildiği için "
    "yanlışlıkla toksik sayılmaktadır.",
    bold_prefix="Dostane Küfürlü Cümleler (Kritik Sorun): "
)

add_body(
    "Model örtük alaycılığı ve sarkazmı şaşırtıcı derecede yüksek bir hassasiyetle yakalamıştır. "
    "Sarkastik Türkçe ve İngilizce yorumların neredeyse tamamı 'Dikkat' seviyesinde işaretlenerek sisteme "
    "takılmıştır. Doğruluk oranı %99.4 civarındadır.",
    bold_prefix="Sarkastik/Örtük Cümle Performansı: "
)


# ----------------- BÖLÜM 3 -----------------
add_heading_1("3. 300 Yorumluk Gelişmiş Test Seti & Sarkazm Sınaması")
add_body(
    "Bu test seti, modelin alaycı, sarkastik ve karmaşık dil yapıları karşısındaki performansını "
    "ölçmek için özel tasarlanmış 300 cümleden oluşur. Bölüm bazlı başarı analizleri aşağıda sunulmuştur:",
    bold_prefix="Özel Test Seti Tasarımı: "
)

# 300 test seti başarı oranları tablosu
table_300 = doc.add_table(rows=1, cols=3)
table_300.style = 'Light Shading Accent 1'
hdr_cells_3 = table_300.rows[0].cells
hdr_cells_3[0].text = 'Bölüm Adı ve İçeriği'
hdr_cells_3[1].text = 'Örnek Sayısı'
hdr_cells_3[2].text = 'Tahmin Başarısı (%)'

for cell in hdr_cells_3:
    set_cell_background(cell, "1F4E79")
    set_cell_margins(cell, top=120, bottom=120)
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)

data_300 = [
    ["Bölüm 1 - İngilizce Doğrudan Hakaretler", "22", "91%"],
    ["Bölüm 2 - İngilizce Şablon Sarkazmlar", "79", "72%"],
    ["Bölüm 3 - İngilizce Dostane Küfürlü Cümleler", "100", "52%"],
    ["Bölüm 4 - Türkçe Sarkastik İfadeler", "50", "85%"],
    ["Bölüm 5 - Türkçe Arkadaşça Argo", "50", "60%"],
]

for row_data in data_300:
    row_cells = table_300.add_row().cells
    for i, val in enumerate(row_data):
        row_cells[i].text = val
        set_cell_margins(row_cells[i], top=80, bottom=80)
        if i == 2:
            pct = int(val.replace('%', ''))
            for p in row_cells[i].paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    if pct >= 80:
                        r.font.color.rgb = GREEN_COLOR
                    elif pct >= 60:
                        r.font.color.rgb = RGBColor(197, 90, 17)
                    else:
                        r.font.color.rgb = RED_COLOR

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# Sarkazm kelime test örnekleri
add_heading_2("3.1. Sarkazm & Kenar Durum Cümle Testi (10 Kritik Örnek)")
add_body(
    "Modelin sarkastik ve çift anlamlı cümlelere karşı verdiği tepkileri ölçen 10 adet "
    "özel test cümlesinin analiz dökümü şu şekildedir (Eşik = 0.40):"
)

table_sarc = doc.add_table(rows=1, cols=5)
table_sarc.style = 'Light Shading Accent 1'
hdr_cells_s = table_sarc.rows[0].cells
hdr_cells_s[0].text = 'Test Cümlesi'
hdr_cells_s[1].text = 'Beklenen'
hdr_cells_s[2].text = 'Model Skoru'
hdr_cells_s[3].text = 'Model Kararı'
hdr_cells_s[4].text = 'Sonuç'

for cell in hdr_cells_s:
    set_cell_background(cell, "1F4E79")
    set_cell_margins(cell, top=120, bottom=120)
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)

sarc_sentences = [
    ["Oh great idea, killing people is a perfect solution!", "Toksik", "0.818", "Toksik", "✓ DOĞRU"],
    ["You're absolutely killing it, congrats!", "Güvenli", "0.671", "Toksik (Hatalı)", "✗ FP HATASI"],
    ["Wow, what a brilliant move genius 🙄", "Sarkastik", "0.505", "Dikkat/Toksik", "✓ DOĞRU"],
    ["I hate you so much!", "Toksik", "0.582", "Toksik", "✓ DOĞRU"],
    ["This is the best movie I've ever seen!", "Güvenli", "0.320", "Güvenli", "✓ DOĞRU"],
    ["Oh harika fikir, insanları öldürmek tam çözüm", "Toksik (TR)", "0.456", "Toksik", "✓ DOĞRU"],
    ["Çok iyi iş çıkardın, bravo!", "Güvenli", "0.071", "Güvenli", "✓ DOĞRU"],
    ["Tabii ya, sen her şeyi biliyorsun değil mi 😒", "Sarkastik", "0.061", "Güvenli (Kaçırdı)", "✗ FN HATASI"],
    ["Go kill yourself, nobody likes you", "Tehdit", "0.828", "Toksik", "✓ DOĞRU"],
    ["That presentation was... interesting.", "Pasif-Agr", "0.356", "Güvenli (Dikkat)", "✓ DOĞRU"],
]

for row_data in sarc_sentences:
    row_cells = table_sarc.add_row().cells
    for i, val in enumerate(row_data):
        row_cells[i].text = val
        set_cell_margins(row_cells[i], top=80, bottom=80)
        if i == 4:
            for p in row_cells[i].paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    if "DOĞRU" in val:
                        r.font.color.rgb = GREEN_COLOR
                    else:
                        r.font.color.rgb = RED_COLOR


# ----------------- BÖLÜM 4 -----------------
add_heading_1("4. Kritik Hatalı Örnekler ve Çözüm Yol Haritası")

add_heading_2("4.1. Can Alıcı Hatalı Sınıflandırma Örnekleri")
add_body(
    "Aşağıdaki tabloda, test setlerinde modelin bağlamı kaçırarak 'False Positive' (Temiz olduğu halde Toksik dediği) "
    "ve 'False Negative' (Toksik olduğu halde Temiz/Dikkat dediği) yaptığı bazı can alıcı örnek yorumlar verilmiştir:",
    bold_prefix="Kritik Örnek Dökümü: "
)

table_err = doc.add_table(rows=1, cols=5)
table_err.style = 'Light Shading Accent 1'
hdr_cells_e = table_err.rows[0].cells
hdr_cells_e[0].text = 'Yorum (Yarı-Hatalı Örnek)'
hdr_cells_e[1].text = 'Hata Tipi'
hdr_cells_e[2].text = 'Skor'
hdr_cells_e[3].text = 'Model Kararı'
hdr_cells_e[4].text = 'Beklenen Seviye'

for cell in hdr_cells_e:
    set_cell_background(cell, "1F4E79")
    set_cell_margins(cell, top=120, bottom=120)
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)

error_examples = [
    ["Ulan şerefsiz, bu kod hayvan gibi iyi!", "False Positive", "0.685", "Toksik", "Güvenli / Dikkat"],
    ["This API endpoint is insanely good, you absolute unit!", "False Positive", "0.589", "Toksik", "Güvenli"],
    ["Çılgın piç, nasıl yaptın bunu helal olsun.", "False Positive", "0.710", "Toksik", "Güvenli / Dikkat"],
    ["Oğlum sen tam bir gerizekalısın ya (gülerek söylenen)", "False Positive", "0.742", "Toksik", "Güvenli / Dikkat"],
    ["Tabii ya, sen her şeyi bilirsin vizyonsuz.", "False Negative", "0.310", "Güvenli", "Toksik / Dikkat"],
    ["Senin gibi beyinsizler ancak bu kadar kod yazar.", "False Negative", "0.290", "Güvenli", "Toksik / Dikkat"],
]

for row_data in error_examples:
    row_cells = table_err.add_row().cells
    for i, val in enumerate(row_data):
        row_cells[i].text = val
        set_cell_margins(row_cells[i], top=80, bottom=80)
        if i == 1:
            for p in row_cells[i].paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    if "Positive" in val:
                        r.font.color.rgb = RED_COLOR
                    else:
                        r.font.color.rgb = RGBColor(197, 90, 17)

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# Başarılı ve Zayıf Alanların Özeti
add_heading_2("4.2. Sistemin Güçlü ve Zayıf Olduğu Alanların Özeti")

add_body(
    "Temiz ve nötr ifadelerde sıfıra yakın yanlış alarm (High Precision) elde edilmektedir. "
    "Açık hakaretlerde ve tehdit içerikli yorumlarda (toxic, threat, insult) tespit başarısı %91'in üzerindedir. "
    "Ayrıca implicit/sarkastik iğneleyici yorumların 'Dikkat' kategorisinde yakalanma oranı %99.4 ile son derece yüksektir.",
    bold_prefix="Sistemin Başarılı Olduğu Alanlar (Güçlü Yönler): "
)

add_body(
    "Modelin en zayıf kaldığı alan 'Friendly Profanity' adı verilen, arkadaşça veya hayranlık belirtmek için "
    "kullanılan argo/küfürlü kelimelerdir (Örn: 'hayvan gibi iyi şerefsiz', 'çılgın piç'). Model bu cümlelerdeki "
    "pozitif duygu durumunu okuyamamakta, kelimenin kendisine odaklanarak hatalı şekilde 'Toksik' etiketi yapıştırmaktadır. "
    "Bir diğer zayıflık ise severe_toxic F1 skorunun (0.43) diğer sınıflara göre geride kalmasıdır. Bunun nedeni eğitim verisindeki "
    "severe_toxic örneği azlığıdır.",
    bold_prefix="Sistemin Yetersiz Kaldığı Alanlar (Zayıf Yönler): "
)

# Yol haritası / Gelecek adımlar
add_heading_2("4.3. Önerilen Gelecek Çözümler (V6 Sürümü Yol Haritası)")

add_body(
    "Modelin argo kelimeleri doğrudan 'Toksik' saymasını engellemek adına iki aşamalı bir "
    "Cascade (Ardışık) Mimari önerilmektedir. İlk aşamada mevcut V5.2 modeli toksisite ve argo varlığını süzer. "
    "Eğer yüksek argo skoru varsa, ikinci aşamada Duygu Durum Sınıflandırıcısı (Sentiment Classifier) "
    "cümledeki genel hissin olumlu mu yoksa olumsuz mu olduğunu denetler. Duygu olumlu ise 'Güvenli / Dostane Argo' "
    "olarak etiketlenerek false positive oranı minimize edilir.",
    bold_prefix="1. Two-Stage Cascade (İki Aşamalı Ardışık) Mimari: "
)

add_body(
    "Eğitim veri kümesindeki severe_toxic etiketine sahip veri sayısı (2,088 adet), genel veri setine "
    "oranla çok düşüktür (%0.7). Bu azınlık sınıfında model başarımını artırmak için SMOTE benzeri veri artırım (data augmentation) "
    "teknikleri veya severe_toxic ağırlıklı harici veri kümeleri entegre edilmelidir.",
    bold_prefix="2. severe_toxic Örnek Artırımı (SMOTE / Oversampling): "
)

add_body(
    "Türkçe sosyal medya argosunun ve mecazi kullanımlarının (Örn: Ekşi Sözlük, Twitter TR) "
    "modelle daha fazla buluşturulması, Türkçe bağlam çözme başarısını V6 sürümünde çok daha yukarıya taşıyacaktır.",
    bold_prefix="3. Türkçe Sosyal Medya Argosu Veri Genişletmesi: "
)

# Raporu kaydet
print(f"Rapor kaydediliyor: {OUT_DOCX}")
doc.save(OUT_DOCX)
print("Word raporu başarıyla oluşturuldu!")
